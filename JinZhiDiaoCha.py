import os
import shutil
import sys
import time
import pythoncom
import win32com.client as wc
from docx import Document


def mkdir(path):
    # 去除首位空格
    path = path.strip()
    # 去除尾部 \ 符号
    path = path.rstrip("\\")

    # 判断路径是否存在
    isExists = os.path.exists(path)
    if not isExists:
        os.makedirs(path)
        return True
    else:
        # 如果目录存在则不创建，并提示目录已存在
        return False

def copy_file (source, target):
    try:
        shutil.copy(source, target)
    except IOError as e:
        print("Unable to copy file. %s" % e)
    except:
        print("Unexpected error:", sys.exc_info())

def read_dir (JinDiao_dict, Guanli_dict, Zizhi_dict, file_path):
    e = os.path.basename(file_path)
    if os.path.isdir(file_path) and not e.startswith('._'):
        son_files = os.listdir(file_path)
        for son_file in son_files:
            son_file_path = file_path + '\\' + son_file
            if os.path.isdir(son_file_path):
                read_dir(JinDiao_dict, Guanli_dict, Zizhi_dict, son_file_path)
            else:
                classify_file(JinDiao_dict, Guanli_dict, Zizhi_dict, son_file_path, son_file)

def classify_file (JinDiao_dict, Guanli_dict, Zizhi_dict, file_path, file_name):
    if not os.path.isdir(file_path) and '._' not in file_path:
        filemt = time.localtime(os.stat(file_path).st_mtime)
        filemt = time.strftime("%Y%m%d%H", filemt)
        if ('尽职调查' in file_name or '尽调' in file_name):
            JinDiao_dict[filemt] = file_path
        elif ('管理人公示' in file_name or '基金业协会' in file_name):
            Guanli_dict[filemt] = file_path
        elif '资质证明' in file_name :
            Zizhi_dict[filemt] = file_path

def move_file (JinDiao_dict, Guanli_dict, Zizhi_dict, target):
    JinDiao_file = target + '尽职调查' + '\\'
    mkdir(JinDiao_file)
    Guanli_dict = sorted(Guanli_dict.items(), key=lambda d:d[0], reverse=True)
    Zizhi_dict = sorted(Zizhi_dict.items(), key=lambda d: d[0], reverse=True)
    for value in JinDiao_dict.values():
        i = 0
        e = os.path.basename(value)
        if '.doc' in e or '.dot' in e or '.rtf' in e:
            copy_file(value, JinDiao_file+str(i)+'-'+e)
            delWordContent(JinDiao_file+str(i)+'-'+e, JinDiao_file+str(i)+'processed-'+e, '评分')
            i += 1

    for value in Guanli_dict:
        e = os.path.basename(value[1])
        copy_file(value[1], target + e)
        break
    for value in Zizhi_dict:
        e = os.path.basename(value[1])
        if '.doc' in e or '.dot' in e or '.rtf' in e:
            copy_file(value[1], target + e)
            text_list = ['投资经理资质证明文件','二、投资经理','二、投资经理资质']
            delWordContent(target + e, target +"-processed-"+ e, text_list)
            break

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None

# 从最后开始替换某字符串几次
def rreplace(s, old, new, occurrence):
    li = s.rsplit(old, occurrence)
    return new.join(li)

def doc_to_docx(doc_name):
    pythoncom.CoInitialize()
    try:
        word = wc.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_name, Encoding='utf-8')
        # 上面的地方只能使用完整绝对地址，相对地址找不到文件，且，只能用“\\”，不能用“/”，哪怕加了 r 也不行，涉及到将反斜杠看成转义字符。

        doc.SaveAs(doc_name.replace('.doc','.docx'), 12, False, "", True, "", False, False, False, False)
        # 转换后的文件,12代表转换后为docx文件
        doc.Close
    except Exception as e:
        print(e)
    finally:
        # 对com操作，一定要确保退出word应用
        if word:
            word.Quit
            del word
        # 释放资源
        pythoncom.CoUninitialize()

def delWordContent(docx_file,dest_file, text_list):
    #读取文本
    if docx_file.endswith('.docx'):
        doc = Document(docx_file)
    else:
        doc_to_docx(docx_file)
        doc = Document(docx_file.replace('.doc', '.docx'))
        print(dest_file)
        dest_file = dest_file.replace('.doc', '.docx')
    paragraphs = doc.paragraphs
    i = 0
    flag = False
    for p in paragraphs:
        i+=1
        #print(str(i))
        print(p.text)
        for text in text_list:
            if p.text.find(text) > -1:
                #print('找到了')
                flag = True
                break
        if flag is True:
            #print('deleting')
            delete_paragraph(p)
    if flag is True:
        #保存为新文件
        doc.save(dest_file)


# 按间距中的绿色按钮以运行脚本。
if __name__ == '__main__':
    filePath = 'F:\\liwidis\\大学\\2022夏季\\2022夏工作文件\\7月\\合规自查-私募尽调\\'
    target = 'F:\\liwidis\\大学\\2022夏季\\2022夏工作文件\\7月\\李博涵整理2\\'
    fileList = os.listdir(filePath)
    file_not_qualified_list = []
    for file in fileList: #各个子基金名字的文件
        e = os.path.basename(file)
        #判断开头是否为‘._’
        if e.startswith('._'):
            continue

        Jindiao = {}  # 是否找到尽职报告
        Guanli = {} # 是否找到管理人公示
        Zizhi = {}  # 是否找到资质证明
        son_file_path = filePath + file
        if os.path.isdir(son_file_path) :
            # grandson_files = os.listdir(son_file_path)
            son_file_target = target + file
            mkdir(son_file_target)
            read_dir(Jindiao, Guanli, Zizhi, son_file_path)
            move_file(Jindiao, Guanli, Zizhi, target+file+'\\')
        if Jindiao and Guanli and Zizhi:
            print(str(file) +'Completed')
        else:
            file_not_qualified_list.append({file:[len(Jindiao), len(Guanli), len(Zizhi)]})

    print(file_not_qualified_list)

# 访问 https://www.jetbrains.com/help/pycharm/ 获取 PyCharm 帮助
